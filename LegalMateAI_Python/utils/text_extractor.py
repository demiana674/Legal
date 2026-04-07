# LegalMateAI_Python/utils/text_extractor.py
"""
Module for extracting text from various document formats
Supported formats: PDF, DOCX, TXT
"""

import os
import PyPDF2
import docx


class TextExtractor:
    """Class for extracting text from different file types"""
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_path (str): Path to PDF file
            
        Returns:
            str: Extracted text from PDF
        """
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num} ---\n"
                        text += page_text + "\n"
                    else:
                        print(f"Warning: Could not extract text from page {page_num}")
        except PyPDF2.errors.PdfReadError as e:
            print(f"PDF read error: {e}")
        except Exception as e:
            print(f"Error extracting PDF: {e}")
        return text

    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path (str): Path to DOCX file
            
        Returns:
            str: Extracted text from DOCX
        """
        text = ""
        try:
            doc = docx.Document(file_path)
            for para_num, para in enumerate(doc.paragraphs, 1):
                if para.text.strip():
                    text += para.text + "\n"
            
            # Also extract text from tables if present
            for table_num, table in enumerate(doc.tables, 1):
                text += f"\n--- Table {table_num} ---\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                        
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
        return text

    @staticmethod
    def extract_from_txt(file_path: str, encoding: str = 'utf-8') -> str:
        """
        Extract text from TXT file
        
        Args:
            file_path (str): Path to TXT file
            encoding (str): File encoding (default: utf-8)
            
        Returns:
            str: Extracted text from TXT
        """
        try:
            # Try multiple encodings if utf-8 fails
            encodings = [encoding, 'utf-8-sig', 'cp1256', 'iso-8859-6']
            for enc in encodings:
                try:
                    with open(file_path, 'r', encoding=enc) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            # If all fail, read as binary and decode with errors='ignore'
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
        except Exception as e:
            print(f"Error extracting TXT: {e}")
            return ""

    @staticmethod
    def extract(file_path: str) -> str:
        """
        Extract text based on file extension
        
        Args:
            file_path (str): Path to file
            
        Returns:
            str: Extracted text from file
        """
        if not os.path.exists(file_path):
            return f"File not found: {file_path}"
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return TextExtractor.extract_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return TextExtractor.extract_from_docx(file_path)
        elif ext == '.txt':
            return TextExtractor.extract_from_txt(file_path)
        else:
            return f"Unsupported file type: {ext}"

    @staticmethod
    def extract_preview(file_path: str, max_chars: int = 500) -> str:
        """
        Extract a preview of the text (first N characters)
        
        Args:
            file_path (str): Path to file
            max_chars (int): Maximum number of characters to return
            
        Returns:
            str: Preview of extracted text
        """
        full_text = TextExtractor.extract(file_path)
        if len(full_text) > max_chars:
            return full_text[:max_chars] + "..."
        return full_text


# ========== Quick Test ==========
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        print(f"Extracting text from: {file_path}")
        print("=" * 50)
        
        text = TextExtractor.extract(file_path)
        print(text)
        print("=" * 50)
        print(f"Total characters: {len(text)}")
    else:
        print("Usage: python text_extractor.py <file_path>")
        print("Example: python text_extractor.py document.pdf")