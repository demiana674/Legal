"""
Text Extractor - Extracts text from various document formats
"""
import os
import PyPDF2
import docx
import logging

logger = logging.getLogger(__name__)

class TextExtractor:
    """
    Extracts text from PDF, DOCX, and TXT files
    """
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- صفحة {page_num} ---\n"
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Could not extract page {page_num}: {e}")
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
        return text
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
        return text
    
    @staticmethod
    def extract_from_txt(file_path: str, encoding: str = 'utf-8') -> str:
        """Extract text from TXT file"""
        encodings = [encoding, 'utf-8-sig', 'cp1256', 'iso-8859-6']
        
        for enc in encodings:
            try:
                with open(file_path, 'r', encoding=enc) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        # Fallback: read with error handling
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            return ""
    
    @staticmethod
    def extract(file_path: str) -> str:
        """Extract text based on file extension"""
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return ""
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return TextExtractor.extract_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return TextExtractor.extract_from_docx(file_path)
        elif ext == '.txt':
            return TextExtractor.extract_from_txt(file_path)
        else:
            logger.warning(f"Unsupported file type: {ext}")
            return f"Unsupported file type: {ext}"